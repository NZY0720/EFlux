from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DEFAULT_OUT = (
    Path(__file__).resolve().parents[2]
    / "artifacts"
    / "docs"
    / "EFlux_Paper_Title_and_Innovations_Bilingual.docx"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 34, 45)
MUTED = RGBColor(92, 104, 116)
LIGHT_FILL = "F4F6F9"
ACCENT_FILL = "E8EEF5"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, *, size=None, bold=None, italic=None, color=None, east_asia="Hiragino Sans GB"):
    # Hiragino Sans GB is installed system-wide on macOS and covers both Latin and
    # Simplified Chinese. Applying it to every OOXML font slot avoids LibreOffice's
    # incomplete East-Asian fallback during the mandatory render check.
    run.font.name = "Hiragino Sans GB"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Hiragino Sans GB")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, side="left", color="2E74B5", size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_label_para(doc, label, text, *, after=7, fill=None, border=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if fill:
        set_paragraph_shading(p, fill)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
    if border:
        set_paragraph_border(p)
        p.paragraph_format.left_indent = Inches(0.12)
    r = p.add_run(label)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_overview_item(doc, number, en_title, zh_title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    n = p.add_run(number)
    set_run_font(n, size=11, bold=True, color=WHITE)
    n._element.get_or_add_rPr().append(_run_shading("2E74B5"))
    t = p.add_run(f"  {en_title}")
    set_run_font(t, size=11, bold=True, color=DARK_BLUE)
    z = p.add_run(f"  |  {zh_title}")
    set_run_font(z, size=10.5, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(7)


def _run_shading(fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def add_innovation(doc, number, en_title, zh_title, zh_summary, en_statement, mechanism, novelty, validation):
    h = doc.add_paragraph(style="Heading 1")
    h.paragraph_format.keep_with_next = True
    r = h.add_run(f"{number}  {en_title}")
    set_run_font(r, size=16, bold=True, color=BLUE)

    zh = doc.add_paragraph(style="Heading 2")
    zh.paragraph_format.keep_with_next = True
    r = zh.add_run(zh_title)
    set_run_font(r, size=13, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph(zh_summary)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for r in p.runs:
        set_run_font(r, size=10.5, color=INK)

    q = doc.add_paragraph()
    q.paragraph_format.space_before = Pt(1)
    q.paragraph_format.space_after = Pt(9)
    q.paragraph_format.left_indent = Inches(0.12)
    q.paragraph_format.right_indent = Inches(0.08)
    q.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(q, LIGHT_FILL)
    set_paragraph_border(q, color="2E74B5", size=20, space=7)
    r = q.add_run("Paper-ready statement  |  英文论文表述\n")
    set_run_font(r, size=9.5, bold=True, color=BLUE)
    r = q.add_run(en_statement)
    set_run_font(r, size=10.25, italic=True, color=INK)

    add_label_para(doc, "Mechanism  |  实现机制：", mechanism)
    add_label_para(doc, "Research novelty  |  创新实质：", novelty)
    add_label_para(doc, "Suggested validation  |  建议验证：", validation, after=12)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Hiragino Sans GB")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Hiragino Sans GB"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    h1._element.rPr.rFonts.set(qn("w:cs"), "Hiragino Sans GB")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Hiragino Sans GB"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    h2._element.rPr.rFonts.set(qn("w:cs"), "Hiragino Sans GB")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("EFlux · Paper Concept Brief")
    set_run_font(r, size=8.5, bold=True, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    add_page_field(fp)


def build(out: Path = DEFAULT_OUT):
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "EFlux Paper Title and Innovations — Bilingual Summary"
    doc.core_properties.subject = "Hierarchical LLM–RL collaboration for VPP trading"
    doc.core_properties.author = "EFlux Research Team"
    doc.core_properties.keywords = "EFlux, LLM, reinforcement learning, VPP, electricity trading"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("PAPER CONCEPT BRIEF  ·  论文构思摘要")
    set_run_font(r, size=9.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    r = title.add_run(
        "EFlux: A Hierarchical LLM–RL Collaboration Framework for Algorithmic Electricity Trading by Virtual Power Plants"
    )
    set_run_font(r, size=20, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("EFlux：面向虚拟电厂算法电力交易的层次化 LLM–RL 协同框架")
    set_run_font(r, size=13, bold=True, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    r = meta.add_run("Research focus  研究方向   ")
    set_run_font(r, size=9.5, bold=True, color=MUTED)
    r = meta.add_run("LLM–RL collaboration · human-guided control · auditable execution · open testbed")
    set_run_font(r, size=9.5, color=INK)

    lead = doc.add_paragraph()
    lead.paragraph_format.left_indent = Inches(0.14)
    lead.paragraph_format.right_indent = Inches(0.10)
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(16)
    lead.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(lead, ACCENT_FILL)
    set_paragraph_border(lead, color="1F4D78", size=24, space=8)
    r = lead.add_run("Core research proposition  |  核心研究命题\n")
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    r = lead.add_run(
        "EFlux studies how strategic reasoning, real-time reinforcement learning, and human supervision can be combined in a physically constrained and fully traceable execution loop for virtual power plant trading.\n"
    )
    set_run_font(r, size=10.5, color=INK)
    r = lead.add_run(
        "EFlux 研究如何将大模型战略推理、强化学习实时执行与人类监督协同起来，并在物理约束和全链路可追溯的条件下完成虚拟电厂算法交易。"
    )
    set_run_font(r, size=10.5, color=INK)

    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run("Contributions at a Glance  |  创新点概览")
    set_run_font(r, size=16, bold=True, color=BLUE)

    add_overview_item(doc, "01", "Hierarchical LLM–RL Collaboration", "层次化 LLM–RL 协同", "A slow LLM strategist provides bounded guidance and meta-control; a fast RL policy executes tactical decisions in real time.")
    add_overview_item(doc, "02", "Human-Guided Semantic Shared Control", "人在环路语义化共享控制", "Operators supervise, steer, or temporarily override the autonomous agent through high-level trading intentions rather than raw orders.")
    add_overview_item(doc, "03", "Physics-Constrained Auditable Execution", "物理约束的可审计执行", "Human, LLM, and RL intentions are compiled into safe orders with provenance across risk checks, delivery, and settlement.")
    add_overview_item(doc, "04", "Open and Reproducible VPP Trading Testbed", "开放可复现的交易测试平台", "Built-in and external agents share one market, physical, risk, delivery, and settlement pipeline for fair benchmarking.")

    doc.add_page_break()

    add_innovation(
        doc,
        "01",
        "Multi-Timescale Hierarchical LLM–RL Collaboration",
        "多时间尺度的层次化 LLM–RL 协同决策",
        "提出双层智能决策架构：低频 LLM 战略层负责市场状态理解、风险姿态与长期目标调节；高频 RL 执行层在结构化动作空间内完成实时战术决策。执行结果持续反馈给战略层，形成“战略指导—战术执行—市场反馈—策略更新”的闭环。",
        "We propose a multi-timescale hierarchical LLM–RL collaboration architecture that separates slow strategic reasoning from fast tactical execution. The LLM produces bounded strategic guidance and learning meta-control, while the RL policy performs real-time trading decisions and returns execution feedback to the strategic layer.",
        "LLM 输出风险预算、SOC 目标、策略偏好、报价偏移和学习元参数；RL 负责每个决策窗口内的快速动作选择；LLM 调用异步运行，不阻塞交易主循环。",
        "创新不在于简单串联 LLM 与 PPO，而在于明确二者的时间尺度、控制权限和反馈接口，使大模型的场景推理能力与轻量策略的实时响应能力形成稳定协作。",
        "对比规则策略、纯 RL、LLM+冻结 RL、LLM+在线 RL；报告收益、失衡、风险拒绝、响应延迟、模型调用成本及不同市场状态下的增益。",
    )

    doc.add_page_break()

    add_innovation(
        doc,
        "02",
        "Human-Guided Semantic Shared Control",
        "人在环路的语义化共享控制机制",
        "提出人—LLM—RL 分层共享控制机制。交易员通过风险预算、目标 SOC、偏好或禁止策略、执行风格、暂停和临时接管等业务语义对自主智能体进行监督与干预，而无需逐笔构造底层订单。干预结束后，控制权可平滑交还给自主策略。",
        "We introduce a human-guided semantic shared-control mechanism that enables operators to supervise, steer, and temporarily override an autonomous LLM–RL agent through bounded high-level trading intentions rather than raw market orders.",
        "支持自主模式、监督模式和接管模式；人类指导经过服务端裁剪、来源标记、持久化和权限校验，并与 LLM 指导和 RL 动作进行统一仲裁。",
        "把人在环路从“人工下单”提升为“战略意图控制”，既保留专业人员对异常行情和风险边界的判断，又保留自动策略的实时执行效率。",
        "设置价格突变、预测失真、流动性枯竭和设备约束等场景，对比完全自主、软指导和临时接管；评估干预次数、响应时间、收益改善、失衡降低及错误指导拦截率。",
    )

    doc.add_page_break()

    add_innovation(
        doc,
        "03",
        "Physics-Constrained Intent Compilation and Auditable Execution",
        "多源意图的物理约束编译与可审计执行",
        "建立从人类指导、LLM 战略意图和 RL 战术动作到订单、成交、物理交付与经济结算的统一执行链。所有决策首先映射为可解释的交易原语，再由确定性编译器生成订单，并通过统一风险网关检查电池 SOC、功率、能源预留、现金和交付承诺。",
        "We develop a physics-constrained intent compilation and auditing mechanism that transforms human, LLM, and RL decisions into executable orders while preserving end-to-end provenance across guidance, action selection, risk validation, market execution, physical delivery, and settlement.",
        "记录指导来源、指导前后动作变化、覆盖关系、订单生成依据、风险接受或拒绝原因、成交结果、SOC、失衡和结算结果；训练与在线部署共享相同的动作语义、编译器和风险入口。",
        "将“安全执行”和“可审计性”统一为一条可重建的意图链，不仅能说明系统做了什么，还能解释最终交易由谁影响、为何执行以及物理与经济结果如何产生。",
        "开展动作编译一致性、约束违规注入、指导冲突仲裁、能量与现金守恒、确定性回放和端到端追踪完整性测试。",
    )

    doc.add_page_break()

    add_innovation(
        doc,
        "04",
        "Open, Extensible, and Reproducible VPP Trading Testbed",
        "开放、可扩展、可复现的虚拟电厂交易实验平台",
        "构建面向研究人员、算法开发者与能源交易人员的标准化实验平台。平台集成异构虚拟电厂资产、P2P 连续双边拍卖与实时电价市场、智能体接入、风险控制、交付结算、回放和多维评测，使外部策略能够在统一条件下开发、测试和比较。",
        "We develop an open, extensible, and reproducible testbed for virtual power plant trading agents. Through a unified agent protocol, the platform exposes the same market observations, physical constraints, risk validation, delivery, and settlement pipeline to both built-in and external strategies, enabling standardized benchmarking and fair comparison.",
        "支持规则策略、传统交易算法、PPO、LLM–RL 与人在环路策略；外部智能体可通过 API、SDK 或统一协议接入；提供固定种子、标准场景、基准算法、隔离配对评测和确定性回放。",
        "平台贡献不只是开放一个网站，而是为不同研究者提供相同执行栈、相同约束和相同评测口径，降低 VPP 交易智能体研究中的复现成本与比较偏差。",
        "发布标准场景和基准结果；验证内置与外部智能体接口等价性、种子复现性、跨策略公平性、压力性能，以及收益、失衡、风险、SOC、延迟和模型成本等指标完整性。",
    )

    doc.add_page_break()

    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run("Paper-Ready Contribution Summary  |  可直接用于论文的贡献总结")
    set_run_font(r, size=16, bold=True, color=BLUE)

    add_label_para(
        doc,
        "English  |  ",
        "EFlux contributes four elements to algorithmic electricity trading by virtual power plants: (1) a multi-timescale hierarchical LLM–RL architecture that separates strategic reasoning from tactical execution; (2) a human-guided semantic shared-control mechanism for bounded supervision and intervention; (3) a physics-constrained intent compilation and auditing pipeline that preserves provenance from guidance to settlement; and (4) an open and reproducible testbed that supports standardized evaluation of built-in and external trading agents.",
        after=14,
        fill=LIGHT_FILL,
        border=True,
    )
    add_label_para(
        doc,
        "中文  |  ",
        "EFlux 面向虚拟电厂算法电力交易形成四项主要贡献：（1）构建分离战略推理与战术执行的多时间尺度层次化 LLM–RL 协同架构；（2）提出支持有界监督、干预与接管的人在环路语义化共享控制机制；（3）建立从多源指导到交付结算的物理约束意图编译与审计链路；（4）提供支持内置与外部交易智能体标准化评测的开放、可扩展、可复现实验平台。",
        after=18,
        fill=LIGHT_FILL,
        border=True,
    )

    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run("Recommended Positioning  |  推荐定位")
    set_run_font(r, size=16, bold=True, color=BLUE)
    add_label_para(
        doc,
        "Primary scientific claim  |  核心科学主张：",
        "EFlux demonstrates how strategic intelligence, tactical learning, and human expertise can be coordinated without allowing any decision source to bypass physical and market safety constraints. / EFlux 展示了如何在不允许任何决策来源绕过物理与市场安全约束的前提下，协同战略智能、战术学习与人类专业知识。",
        after=10,
    )
    add_label_para(
        doc,
        "System contribution  |  系统贡献：",
        "The same execution and evaluation stack is exposed to researchers and external agents, making the proposed methods testable, comparable, and reproducible. / 同一套执行与评测链路向研究人员及外部智能体开放，使所提出的方法能够被测试、比较和复现。",
        after=10,
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(out)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the bilingual EFlux paper brief.")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    build(args.out)


if __name__ == "__main__":
    main()
