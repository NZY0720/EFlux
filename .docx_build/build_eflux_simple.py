from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/nikoneo/Desktop/EFlux_Paper_Title_and_Innovations_Bilingual.docx")
FONT = "Hiragino Sans GB"
INK = RGBColor(28, 36, 45)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)


def font(run, *, size=11, bold=False, italic=False, color=INK):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rpr.rFonts.set(qn(f"w:{key}"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    font(r, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_body(doc, label, text, *, english=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.10
    lr = p.add_run(f"{label}  ")
    font(lr, size=10.5, bold=True, color=DARK_BLUE)
    tr = p.add_run(text)
    font(tr, size=10.5, italic=english, color=INK)


def add_innovation(doc, number, en_title, zh_title, zh_text, en_text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True
    nr = h.add_run(f"{number}  ")
    font(nr, size=14, bold=True, color=BLUE)
    er = h.add_run(en_title)
    font(er, size=14, bold=True, color=BLUE)

    z = doc.add_paragraph()
    z.paragraph_format.space_before = Pt(0)
    z.paragraph_format.space_after = Pt(7)
    z.paragraph_format.keep_with_next = True
    zr = z.add_run(zh_title)
    font(zr, size=12, bold=True, color=DARK_BLUE)

    add_body(doc, "中文", zh_text)
    add_body(doc, "English", en_text, english=True)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    page_number(section.footer.paragraphs[0])

    title_label = doc.add_paragraph()
    title_label.paragraph_format.space_before = Pt(0)
    title_label.paragraph_format.space_after = Pt(6)
    r = title_label.add_run("论文标题  |  Paper Title")
    font(r, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.keep_with_next = True
    r = title.add_run(
        "EFlux: A Hierarchical LLM–RL Collaboration Framework for Algorithmic Electricity Trading by Virtual Power Plants"
    )
    font(r, size=18, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("EFlux：面向虚拟电厂算法电力交易的层次化 LLM–RL 协同框架")
    font(r, size=12.5, bold=True, color=DARK_BLUE)

    add_innovation(
        doc,
        "01",
        "Multi-Timescale Hierarchical LLM–RL Collaboration",
        "多时间尺度的层次化 LLM–RL 协同决策",
        "提出由低频 LLM 战略层与高频 RL 执行层组成的协同架构。LLM 负责市场状态理解、风险姿态和长期目标调节，RL 负责结构化动作空间内的实时交易决策，并将执行结果反馈给战略层。",
        "We propose a multi-timescale hierarchical LLM–RL architecture in which a slow LLM strategist performs market reasoning and strategic control, while a fast RL policy makes real-time tactical decisions and returns execution feedback to the strategic layer.",
    )

    add_innovation(
        doc,
        "02",
        "Human-Guided Semantic Shared Control",
        "人在环路的语义化共享控制机制",
        "交易员可通过风险预算、目标 SOC、策略偏好、暂停和临时接管等高层语义对智能体进行监督与干预，而无需逐笔构造底层订单；干预结束后，控制权可交还给自主策略。",
        "We introduce a human-guided semantic shared-control mechanism that enables operators to supervise, steer, and temporarily override the autonomous agent through bounded high-level trading intentions rather than raw market orders.",
    )

    doc.add_page_break()

    add_innovation(
        doc,
        "03",
        "Physics-Constrained Intent Compilation and Auditable Execution",
        "多源意图的物理约束编译与可审计执行",
        "将人类指导、LLM 战略意图和 RL 战术动作统一映射为可解释的交易原语，再经确定性编译和风险网关生成满足 SOC、功率、能源、现金及交付约束的订单，并完整记录从指导到结算的决策来源。",
        "We develop a physics-constrained intent compilation and auditing mechanism that transforms human, LLM, and RL decisions into safe executable orders while preserving end-to-end provenance from guidance and risk validation to delivery and settlement.",
    )

    add_innovation(
        doc,
        "04",
        "Open, Extensible, and Reproducible VPP Trading Testbed",
        "开放、可扩展、可复现的虚拟电厂交易测试平台",
        "提供面向研究人员和开发者的统一测试平台。内置与外部智能体可通过 API、SDK 或智能体协议接入同一市场、物理、风控、交付和结算链路，并利用标准场景、基准策略和确定性回放进行公平比较。",
        "We provide an open, extensible, and reproducible testbed in which built-in and external agents access the same market, physical, risk, delivery, and settlement pipeline through a unified protocol, enabling standardized testing and fair comparison.",
    )

    doc.core_properties.title = "EFlux Paper Title and Four Innovations — Bilingual"
    doc.core_properties.author = "EFlux Research Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
